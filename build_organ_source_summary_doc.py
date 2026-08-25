import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(r"D:\undergra\na_nn\UNI\outputs\multiclass")
SUMMARY_PATH = OUTPUT_DIR / "organ_source_summary.json"
IMAGE_PRED_PATH = OUTPUT_DIR / "image_predictions.csv"
GROUP_PRED_PATH = OUTPUT_DIR / "group_predictions.csv"
SKIPPED_PATH = OUTPUT_DIR / "skipped_images.csv"
DOCX_OUT = Path(r"C:\Users\y0113\Documents\大创\UNI_organ_source_team_summary.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r)
    return p


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(text))
        set_font(r, size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            if i > 0 and str(text).replace(".", "", 1).replace("%", "").isdigit():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_font(r, size=font_size)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def pct(x):
    return f"{x * 100:.2f}%"


def load_csv(path):
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def main():
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    image_rows = load_csv(IMAGE_PRED_PATH)
    group_rows = load_csv(GROUP_PRED_PATH)
    skipped_rows = load_csv(SKIPPED_PATH)
    wrong_images = [r for r in image_rows if str(r.get("correct", "")).lower() == "false"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("UNI 组织来源分类工作总结（团队共享版）")
    set_font(r, size=20, bold=True, color="0B2545")

    sub = doc.add_paragraph()
    r = sub.add_run("基于显微镜图像 patch、冻结 UNI 特征和 Logistic Regression 的四分类 baseline")
    set_font(r, size=11, color="555555")

    doc.add_heading("阅读导引", level=1)
    add_para(
        doc,
        "本文档用于团队内部同步目前显微镜图像分类工作的阶段性进展。建议先阅读“工作目标”和“当前结论”，再查看数据构成、方法流程和主要结果。"
        "需要复现实验的同学重点查看“方法流程”和“交付物位置”；需要判断下一步研究方向的同学重点查看“对结果的解释”和“尚未完成与后续方向”。"
    )
    add_table(
        doc,
        ["项目", "当前状态"],
        [
            ["任务类型", "组织来源四分类 baseline"],
            ["输入数据", "肝原位、皮下瘤、类器官、扁桃体显微镜图像"],
            ["核心方法", "UNI 冻结特征提取 + Logistic Regression 分类器"],
            ["评估策略", "按 label + slide_id 做样本组留出，降低同一样片泄露"],
            ["当前用途", "验证来源分类可行性；不是病变区域识别模型"],
        ],
        widths=[1.7, 4.5],
    )

    doc.add_heading("1. 工作目标", level=1)
    add_para(
        doc,
        "本阶段目标是建立一个可复现的组织来源分类 baseline。输入为显微镜图像，输出为四类来源标签："
        "liver_orthotopic、tumor_subcutaneous、organoid_crf、tonsil。该任务用于验证 UNI 预训练特征能否区分不同来源/组织域，"
        "不是直接识别病变区域，也不是判断 marker 阳性或阴性。"
    )
    add_para(
        doc,
        "任务边界：当前标签来自整张图像所在来源，而不是医生标注的病灶区域。因此，本阶段结果不能直接用于“图像中哪里是病变区域”的定位。"
    )

    doc.add_heading("2. 已完成工作", level=1)
    for item in [
        "整理并统一了肝原位与皮下瘤数据的类别目录和命名规则。",
        "从 CRF/CRF2 数据中额外整理出 organoid_crf 和 tonsil 两类；TA 系列暂不纳入。",
        "在服务器 GPU 环境中完成四分类推理与评估流程。",
        "采用冻结 UNI 模型提取 1024 维图像 patch 特征，并用 Logistic Regression 训练分类器。",
        "使用按样本组留出的评估方式，降低同一样片、不同视野、不同倍率或不同 marker 带来的数据泄露风险。",
        "生成 selected_images、patch_manifest、patch_predictions、image_predictions、group_predictions 和 summary 等输出文件。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. 交付物位置", level=1)
    add_table(
        doc,
        ["文件", "用途"],
        [
            ["organ_source_summary.json", "汇总指标、混淆矩阵、类别数量、异常图像和分层结果"],
            ["image_predictions.csv", "每张图像的真实类别、预测类别、是否正确、patch 投票结果"],
            ["group_predictions.csv", "每个样本组的真实类别、预测类别、marker/倍率构成"],
            ["patch_predictions.csv", "每个 patch 的预测结果，适合排查具体图像局部错误"],
            ["patch_manifest.csv", "patch 来源、坐标、tissue score 与图像元信息"],
            ["uni_features.npy", "UNI 提取的 1024 维 patch embedding"],
        ],
        widths=[2.0, 4.2],
    )

    doc.add_heading("4. 数据构成", level=1)
    label_counts = summary["label_counts_image"]
    patch_counts = summary["label_counts_patch"]
    data_rows = []
    for label in ["liver_orthotopic", "tumor_subcutaneous", "organoid_crf", "tonsil"]:
        data_rows.append([label, label_counts.get(label, 0), patch_counts.get(label, 0)])
    add_table(doc, ["类别", "图像数", "patch 数"], data_rows, widths=[2.7, 1.7, 1.7])
    add_para(doc, f"总图像数为 {summary['num_images']}，总 patch 数为 {summary['num_patches']}，UNI 特征矩阵形状为 {summary['features_shape']}。")

    doc.add_heading("5. 方法流程", level=1)
    for item in [
        "每张原始图像先按网格切成 512 x 512 patch；若图像小于 patch 尺寸，则使用整张图像。",
        "对每个 patch 计算简易 tissue score，优先选择纹理和组织区域更明显的 patch。",
        "每张图默认最多取 16 个 patch，不保存 patch 图像，只保存特征与 manifest。",
        "使用 UNI 原版权重，模型作为冻结特征提取器，不进行端到端深度学习微调。",
        "每个 patch 经过 UNI 后得到 1024 维 embedding，再用 StandardScaler + Logistic Regression 训练分类器。",
        "评估采用 group holdout：每次留出一个 label + slide_id 组作为测试集，其余组作为训练集。",
    ]:
        add_number(doc, item)

    doc.add_heading("6. 主要结果", level=1)
    result_rows = [
        ["Patch level", pct(summary["patch"]["accuracy"]), pct(summary["patch"]["balanced_accuracy"])],
        ["Image level", pct(summary["image"]["accuracy"]), pct(summary["image"]["balanced_accuracy"])],
        ["Group level", pct(summary["group"]["accuracy"]), pct(summary["group"]["balanced_accuracy"])],
    ]
    add_table(doc, ["评估层级", "Accuracy", "Balanced accuracy"], result_rows, widths=[2.4, 1.8, 2.2])

    image_cm = summary["image"]["confusion_matrix"]
    labels = summary["image"]["confusion_matrix_labels"]
    cm_rows = []
    for label, row in zip(labels, image_cm):
        cm_rows.append([label] + row)
    add_table(doc, ["真实类别 \\ 预测类别"] + labels, cm_rows, widths=[1.6, 1.15, 1.3, 1.15, 0.95])

    doc.add_heading("7. 错误样本与异常图像", level=1)
    if wrong_images:
        wrong_rows = [
            [
                r["file_name"],
                r["true_label"],
                r["pred_label"],
                r["group_key"],
                r["magnification"],
                r["votes"],
            ]
            for r in wrong_images
        ]
        add_table(doc, ["文件名", "真实类别", "预测类别", "group", "倍率", "投票"], wrong_rows, widths=[2.0, 1.0, 1.0, 1.3, 0.5, 0.7], font_size=8.5)
        add_para(doc, "两张错误图像均来自 organoid2 的 5x 图像，均被预测为 tonsil，说明低倍类器官图像与扁桃体之间存在局部特征混淆。")
    else:
        add_para(doc, "image-level 未发现错误样本。")

    if skipped_rows:
        skipped_table_rows = [[Path(r["source_path"]).name, r["error"]] for r in skipped_rows]
        add_table(doc, ["被跳过图像", "错误原因"], skipped_table_rows, widths=[4.5, 1.6], font_size=8.5)
    add_para(doc, "被跳过图像没有参与最终预测，因此图像级有效评估样本数为 462 张。")

    doc.add_heading("8. 对结果的解释", level=1)
    for item in [
        "当前结果说明 UNI 特征对组织来源/数据域差异非常敏感，能够较好地区分肝原位、皮下瘤、类器官和扁桃体。",
        "高准确率不能直接解释为模型已经学会识别病变区域，因为训练标签是整张图像的来源标签，而不是区域级病变标注。",
        "marker 与类别仍然存在明显相关性，例如 MPA_P4 和 TCR_REP 只出现在 liver_orthotopic，MFAP4/TCR_REF/TCR_RFP 只出现在 tumor_subcutaneous。模型可能同时利用了组织形态、染色风格、制片差异和拍摄条件。",
        "group-level 结果为 100%，但 organoid_crf 和 tonsil 的独立样片组数量较少，统计说服力有限。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9. 尚未完成与后续方向", level=1)
    for item in [
        "尚未进行端到端深度学习微调；当前只是冻结 UNI 特征 + 传统分类器。",
        "尚未建立病变区域级标注，因此不能训练真正的病灶定位或区域分割模型。",
        "尚未完成跨 marker 泛化、跨倍率泛化和跨实验批次泛化的系统评估。",
        "尚未加入 embedding 可视化；建议用 PCA、t-SNE 或 UMAP 同时标出类别、marker 和倍率。",
        "建议单独分析 organoid_crf vs tonsil，尤其检查 organoid2 的 5x 图像为何被预测为 tonsil。",
        "后续如果加入 class_summary 中更多类别，应优先明确每个任务的生物学问题，不宜把所有杂类混成一个大任务。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("10. 队友可接手的具体工作", level=1)
    for item in [
        "检查 3 张被跳过图像能否重新导出或重新截图，尤其是 PNG truncated 与 TIFF decoder error。",
        "人工查看 organoid2 的 5x 错误样本，判断是否存在背景、低倍视野或图像质量导致的混淆。",
        "补充 organoid_crf 和 tonsil 的样本量；当前两类独立样片组较少，group-level 100% 的统计稳定性有限。",
        "做 embedding 可视化，把类别、marker、倍率同时标出来，判断模型是否按组织来源分群，还是被染色或倍率驱动。",
        "如要进入病变区域识别，需要先建立区域级标注或至少图像级弱标注方案。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("11. 当前结论", level=1)
    add_para(
        doc,
        "本阶段已经完成了一个可运行、可复现、基于样本组留出的四分类 baseline。结果显示，UNI 预训练特征能够很好地区分当前数据中的组织来源。"
        "下一阶段的关键不是继续追求更高的来源分类准确率，而是验证模型依赖的特征来源，并逐步从来源分类过渡到更接近实际需求的病变区域识别。"
    )

    doc.core_properties.title = "UNI 组织来源分类工作总结（团队共享版）"
    doc.core_properties.subject = "显微镜图像组织来源分类 baseline 总结"
    doc.core_properties.author = "Codex"
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
